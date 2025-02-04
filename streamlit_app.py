import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from deep_translator import GoogleTranslator

def translate_doc(doc, destination='hi'):
    """
    Translate a Word document while preserving formatting and structure.
    :param doc: Word doc object (from Document class)
    :param destination: Target language (default is Hindi 'hi')
    """
    translator = GoogleTranslator(source='auto', target=destination)
    
    # Translate paragraphs while preserving formatting
    for p in doc.paragraphs:
        if p.text.strip():
            try:
                translated_text = translator.translate(p.text.strip()) or p.text
                for run in p.runs:
                    run.text = ''
                new_run = p.add_run(translated_text)
                new_run.font.size = Pt(9)
            except Exception as e:
                print(f"Error translating paragraph: {e}")
    
    # Translate table cells while preserving structure
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    try:
                        full_text = "\n".join([para.text.strip() for para in cell.paragraphs])
                        translated_text = translator.translate(full_text) or full_text
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.text = ''
                        cell.text = ''
                        new_para = cell.paragraphs[0].add_run(translated_text)
                        new_para.font.size = Pt(9)
                    except Exception as e:
                        print(f"Error translating cell text: {e}")
    
    return doc

def main():
    st.title("Word Document Translator")
    
    uploaded_file = st.file_uploader("Upload a Word Document", type=["docx"])
    if uploaded_file:
        doc = Document(uploaded_file)
        
        language_options = {
            "Bengali": "bn", "Hindi": "hi", "Odia": "or", "Punjabi": "pa", 
            "Tamil": "ta", "Telegu": "te", "Gujarati": "gu", "Malayalam": "ml"
        }
        target_language = st.selectbox("Select Target Language", options=list(language_options.keys()))
        language_code = language_options[target_language]
        
        if st.button("Translate Document"):
            with st.spinner('Translating...'):
                translated_doc = translate_doc(doc, language_code)
                
                with open("translated_document.docx", "wb") as f:
                    translated_doc.save(f)
                
                with open("translated_document.docx", "rb") as f:
                    st.download_button(
                        label="Download Translated Document",
                        data=f,
                        file_name="translated_document.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                st.success("Translation complete!")

if __name__ == '__main__':
    main()
